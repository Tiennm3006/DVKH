import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches

# ===== HÀM TIỆN ÍCH =====
def create_bar_chart(df, col_name, title):
    fig, ax = plt.subplots()
    df_sorted = df.sort_values(by=col_name, ascending=False)
    bars = ax.bar(df_sorted[df.columns[1]], df_sorted[col_name])
    ax.set_title(title)
    ax.set_ylabel(col_name)
    plt.xticks(rotation=45, ha='right')

    for bar in bars:
        height = bar.get_height()
        ax.annotate(f'{height:.2f}%',
                    xy=(bar.get_x() + bar.get_width() / 2, height),
                    xytext=(0, 3),
                    textcoords="offset points",
                    ha='center', va='bottom', fontsize=8)

    fig.tight_layout()
    buffer = BytesIO()
    plt.savefig(buffer, format="png")
    plt.close(fig)
    buffer.seek(0)
    return buffer

def generate_report(df, col_label, chart_all, chart_top, chart_bot, top3, bot3, filename):
    doc = Document()
    doc.add_heading('BÁO CÁO PHÂN TÍCH', 0)

    doc.add_heading('Bảng dữ liệu tổng:', level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    doc.add_paragraph('Biểu đồ kết quả thực hiện tổng:')
    doc.add_picture(chart_all, width=Inches(5.5))

    doc.add_paragraph('Top 3 cao nhất:')
    table_top = doc.add_table(rows=1, cols=len(top3.columns))
    table_top.style = 'Table Grid'
    for i, col in enumerate(top3.columns):
        table_top.rows[0].cells[i].text = str(col)
    for _, row in top3.iterrows():
        row_cells = table_top.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    doc.add_picture(chart_top, width=Inches(5.5))

    doc.add_paragraph('Bottom 3 thấp nhất:')
    table_bot = doc.add_table(rows=1, cols=len(bot3.columns))
    table_bot.style = 'Table Grid'
    for i, col in enumerate(bot3.columns):
        table_bot.rows[0].cells[i].text = str(col)
    for _, row in bot3.iterrows():
        row_cells = table_bot.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    doc.add_picture(chart_bot, width=Inches(5.5))

    doc.save(filename)
    return filename

# ===== GIAO DIỆN STREAMLIT =====
st.set_page_config(layout="wide")
st.title("Phân tích tỷ lệ CSKH và yêu cầu KH")

tabs = st.tabs(["Tỷ lệ phát triển App CSKH", "Tỷ lệ đúng hạn xử lý yêu cầu KH"])

with tabs[0]:
    st.header("Phân tích tỷ lệ phát triển App CSKH")
    uploaded_file1 = st.file_uploader("Tải lên file Excel App CSKH", type="xlsx", key="app")
    if uploaded_file1:
        df1 = pd.read_excel(uploaded_file1, skiprows=2)
        df1.columns = ['STT', 'Điện lực', 'Số lượng KH quản lý', 'Số lượng thực hiện App', 'Tỷ lệ thực hiện qua App']
        df1 = df1.dropna(subset=['Điện lực'])
        df1['Tỷ lệ thực hiện qua App (%)'] = (df1['Tỷ lệ thực hiện qua App'] * 100).round(2)
        # Tách 'Công ty' ra nếu có
        df1_company = df1[df1['Điện lực'].str.lower().str.contains('công ty', na=False)]
        df1 = df1[~df1['Điện lực'].str.lower().str.contains('công ty', na=False)]
        df1 = df1.sort_values(by='Tỷ lệ thực hiện qua App (%)', ascending=False)
        df1 = pd.concat([df1, df1_company], ignore_index=True)

        selected_unit = st.selectbox("Chọn Điện lực để lọc (hoặc để trống xem toàn bộ):", ["-- Tất cả --"] + df1['Điện lực'].unique().tolist())
        if selected_unit != "-- Tất cả --":
            df1 = df1[df1['Điện lực'] == selected_unit]

        st.dataframe(df1)

        top3 = df1[~df1['Điện lực'].str.lower().str.contains('công ty', na=False)].head(3)
        bot3 = df1[~df1['Điện lực'].str.lower().str.contains('công ty', na=False)].tail(3).sort_values(by='Tỷ lệ thực hiện qua App (%)')

        chart_all = create_bar_chart(df1[~df1['Điện lực'].str.lower().str.contains('công ty', na=False)], 'Tỷ lệ thực hiện qua App (%)', 'Biểu đồ tổng thể')
        chart_top = create_bar_chart(top3, 'Tỷ lệ thực hiện qua App (%)', 'Top 3 Điện lực cao nhất')
        chart_bot = create_bar_chart(bot3, 'Tỷ lệ thực hiện qua App (%)', 'Bottom 3 Điện lực thấp nhất')

        st.image(chart_all)
        st.image(chart_top)
        st.image(chart_bot)

        if st.button("Xuất báo cáo Word", key="report1"):
            file_path = "Bao_cao_App_CSKH.docx"
            generate_report(df1, 'Tỷ lệ thực hiện qua App (%)', chart_all, chart_top, chart_bot, top3, bot3, file_path)
            with open(file_path, "rb") as f:
                st.download_button("Tải báo cáo", f, file_name=file_path)

with tabs[1]:
    st.header("Phân tích tỷ lệ đúng hạn xử lý yêu cầu KH")
    uploaded_file2 = st.file_uploader("Tải lên file Excel Yêu cầu KH", type="xlsx", key="request")
    if uploaded_file2:
        df2 = pd.read_excel(uploaded_file2, skiprows=3)
        df2.columns = ['STT', 'Đơn vị', 'Số yêu cầu xử lý', 'Phiếu trễ hạn', 'Tỷ lệ trễ hạn',
                       'Phiếu đúng hạn', 'Tỷ lệ đúng hạn', 'Kế hoạch', 'So sánh']
        df2 = df2.dropna(subset=['Đơn vị'])
        df2 = df2[['STT', 'Đơn vị', 'Số yêu cầu xử lý', 'Phiếu trễ hạn', 'Tỷ lệ trễ hạn']]
        df2['Tỷ lệ trễ hạn (%)'] = (df2['Phiếu trễ hạn'] / df2['Số yêu cầu xử lý'] * 100).round(2)
        # Tách 'Công ty' ra nếu có
        df2_company = df2[df2['Đơn vị'].str.lower().str.contains('công ty', na=False)]
        df2 = df2[~df2['Đơn vị'].str.lower().str.contains('công ty', na=False)]
        df2 = df2.sort_values(by='Tỷ lệ trễ hạn (%)', ascending=True)
        df2 = pd.concat([df2, df2_company], ignore_index=True)

        selected_dv = st.selectbox("Chọn Đơn vị để lọc (hoặc để trống xem toàn bộ):", ["-- Tất cả --"] + df2['Đơn vị'].unique().tolist())
        if selected_dv != "-- Tất cả --":
            df2 = df2[df2['Đơn vị'] == selected_dv]

        st.dataframe(df2)

        top3 = df2[~df2['Đơn vị'].str.lower().str.contains('công ty', na=False)].tail(3).sort_values(by='Tỷ lệ trễ hạn (%)', ascending=False)
        bot3 = df2[~df2['Đơn vị'].str.lower().str.contains('công ty', na=False)].head(3)

        chart_all = create_bar_chart(df2[~df2['Đơn vị'].str.lower().str.contains('công ty', na=False)], 'Tỷ lệ trễ hạn (%)', 'Biểu đồ tổng thể')
        chart_top = create_bar_chart(top3, 'Tỷ lệ trễ hạn (%)', 'Top 3 Đơn vị trễ hạn cao nhất')
        chart_bot = create_bar_chart(bot3, 'Tỷ lệ trễ hạn (%)', 'Bottom 3 Đơn vị trễ hạn thấp nhất')

        st.image(chart_all)
        st.image(chart_top)
        st.image(chart_bot)

        if st.button("Xuất báo cáo Word", key="report2"):
            file_path = "Bao_cao_Yeu_cau_KH.docx"
            generate_report(df2, 'Tỷ lệ trễ hạn (%)', chart_all, chart_top, chart_bot, top3, bot3, file_path)
            with open(file_path, "rb") as f:
                st.download_button("Tải báo cáo", f, file_name=file_path)
